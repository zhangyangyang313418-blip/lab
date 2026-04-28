from __future__ import annotations

import json
import re
import sys
import xml.etree.ElementTree as ET
import zipfile
from collections import defaultdict
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader


SOURCE_ROOT = Path("/Users/clytia/Desktop/标准库/客户标准/捷豹路虎JLR")
OUTPUT_DIR = Path("/Users/clytia/Desktop/Codex/产品测试流程自动化/outputs/jlr-standards")
OUTPUT_JSON = OUTPUT_DIR / "jlr_standards_extracted.json"
OUTPUT_MD = OUTPUT_DIR / "JLR标准名称版本号整理.md"

CATEGORY_ORDER = {
    "环境可靠性测试": 0,
    "噪声测试": 1,
    "材料性能测试": 2,
    "VOC/气味/禁限物质": 3,
    "EMC/电磁兼容": 4,
    "光学/HUD/显示性能": 5,
    "电气/电子设计": 6,
    "电气/电子测试": 7,
    "软件/通信/网络安全": 8,
    "CAD/工程图纸": 9,
    "HMI/开关/座舱": 10,
    "外观/感知质量": 11,
    "整车/装配/试验准备": 12,
    "结构/连接器/NVH": 13,
    "管理/流程/通用规范": 14,
    "其他参考": 99,
}

TITLE_ZH_BY_CODE = {
    "2048-EE": "环境试验选择矩阵",
    "JLR-HTG-564459": "环境试验选择矩阵",
    "STJLR.00.5180": "整车耐腐蚀性",
    "STJLR.00.5235": "整车箱体腐蚀验证要求",
    "STJLR.00.7001": "车辆耐久寿命",
    "STJLR.01.5202": "影响车内空气质量的车身部件环境质量要求",
    "STJLR.15.5031": "驾驶员信息系统环境工作要求",
    "STJLR.18.233": "电气和电子部件环境要求",
    "TPJLR.00.029": "标准车辆噪声与振动测量位置",
    "TPJLR.00.065": "高温气候评价",
    "TPJLR.00.141": "全球12周整车加速腐蚀试验程序",
    "TPJLR.00.165": "高日照负荷下座舱电气及MHEV系统热试验程序",
    "TPJLR.00.170": "温带气候评价",
    "TPJLR.00.187": "部件或系统异响振动台试验",
    "TPJLR.01.126": "仪表板和中控台总成扩展耐久试验",
    "TPJLR.15.045": "聚合物显示屏环境应力开裂试验",
    "TPJLR.18.125": "电气和电子部件环境兼容性试验",
    "STJLR.01.5019": "未喷漆塑料件外观质量",
    "STJLR.01.5020": "内饰喷漆件外观质量",
    "STJLR.01.5023": "内饰饰件、模塑件或模塑喷漆件可见分缝",
    "STJLR.01.5026": "外饰喷漆件外观质量",
    "STJLR.01.5195": "注塑成型部件Moldflow分析",
    "STJLR.50.5071": "铝、镁或铁铸件制造的车辆部件",
    "STJLR.50.5151": "镀铬部件性能要求",
    "STJLR.50.5353": "无铅焊料性能要求",
    "STJLR.51.5006": "内饰应用中胶粘剂性能表征",
    "STJLR.51.5181": "自粘胶带系统性能要求",
    "STJLR.51.5182": "标签和贴花性能要求",
    "STJLR.51.5204": "皮革材料性能",
    "STJLR.51.5205": "含胶原材料性能",
    "STJLR.51.5223": "内饰临时保护系统",
    "STJLR.51.5224": "包含柔性表皮材料的仪表板总成性能",
    "STJLR.51.5225": "内饰总成性能规范",
    "STJLR.51.5226": "生产材料标准要求",
    "STJLR.51.5227": "黑盒/灰盒件材料控制",
    "STJLR.51.5228": "缝纫线材料性能",
    "STJLR.51.5229": "内饰材料和部件环境质量要求",
    "STJLR.51.5232": "内饰材料TÜV Toxproof认证要求",
    "STJLR.51.5242": "内饰油漆和清漆性能",
    "STJLR.51.5262": "非装饰塑料件性能要求",
    "STJLR.51.5270": "聚合物中再生材料的使用",
    "STJLR.51.5281": "聚合物薄膜材料性能",
    "STJLR.51.5282": "非化石来源聚合物薄膜材料性能",
    "STJLR.51.5301": "橡胶件性能要求",
    "STJLR.51.5306": "热塑性弹性体（TPE）性能要求",
    "STJLR.51.5341": "地毯材料性能",
    "STJLR.51.5361": "轻载织物材料性能",
    "STJLR.51.5362": "重载织物材料性能",
    "STJLR.51.5363": "超细纤维织物材料性能",
    "STJLR.51.5366": "内饰SVO附件用纺织品材料性能",
    "STJLR.51.5381": "涂漆装饰内饰件（木饰面）",
    "STJLR.99.9999": "限用物质管理标准",
    "TPJLR.01.425": "注塑成型部件Moldflow分析",
    "TPJLR.52.009": "微划伤性能评估：线性磨耗法",
    "TPJLR.52.010": "损伤性能评估：光泽评估方法",
    "TPJLR.52.061": "油漆附着力试验方法",
    "TPJLR.52.104": "内饰材料VOC释放量测定",
    "TPJLR.52.109": "整车内饰VOC和醛类释放量测定",
    "TPJLR.52.155": "汽车内饰液体斑渍试验",
    "TPJLR.52.156": "内饰耐汗液性：浸泡试验",
    "TPJLR.52.165": "内饰材料耐护肤霜、驱虫剂和车载插电香氛",
    "TPJLR.52.209": "白色无背衬乙烯基迁移污染耐受性",
    "TPJLR.52.210": "内饰材料可清洁性试验",
    "TPJLR.52.351": "耐湿性：通用要求",
    "TPJLR.52.352": "耐热老化：通用要求",
    "TPJLR.52.353": "加速环境老化",
    "TPJLR.52.456": "有机涂层刮擦附着力",
    "TPJLR.52.458": "内饰材料、部件和总成气味测定与评估",
    "TPJLR.52.656": "内饰喷漆表面水解试验",
    "JLR-EMC-CP": "电气/电子部件和子系统EMC符合性程序",
    "STJLR.18.084": "整车EMC标准",
    "STJLR.18.209": "全屏蔽电源线高压交流或直流部件EMC要求",
    "CADJLR.01.01": "技术产品文档：2D和3D CAD定义及修订",
    "CADJLR.01.02": "技术产品文档：字体、缩写和符号",
    "CADJLR.02.01": "工程内容",
    "CADJLR.02.02": "工程内容：品牌标识和编码",
    "CADJLR.02.03": "工程内容：标准注释",
    "CADJLR.03.01": "数字CAD和3D模型要求",
    "CADJLR.04.01": "几何产品规范（GPS）和尺寸标注",
    "STJLR.00.5104": "人机界面（HMI）设计策略：通用色度颜色标准",
    "STJLR.01.5018": "整车视野范围",
    "STJLR.13.5007": "抬头显示风挡玻璃要求",
    "STJLR.13.5009": "抬头显示光学性能标准",
    "STJLR.13.5022": "抬头显示（HUD）布置要求",
    "STJLR.15.5018": "显示均匀性标准",
    "STJLR.15.5019": "显示屏本体最大畸变",
    "STJLR.15.5020": "防反射涂层颜色和反射率",
    "STJLR.15.5023": "显示器Zenbara标准",
    "STJLR.15.5027": "显示器和开关触控标准",
    "STJLR.15.5028": "通用电子显示性能标准",
    "STJLR.15.5036": "显示部件颗粒暴露要求",
    "TPJLR.00.168": "通用电子显示性能试验程序",
    "TPJLR.13.001": "抬头显示光学性能评估",
    "TPJLR.13.002": "整车级抬头显示光学性能评估",
    "TPJLR.13.004": "抬头显示阳光可读性对策验证",
    "TPJLR.15.035": "显示均匀性测量试验程序",
    "TPJLR.15.051": "显示部件颗粒暴露试验程序",
    "STJLR.05.5003": "传动系统控制模块标定和仪表通信协议标准",
    "STJLR.07.5003": "TCM-AZG诊断工具接口工程标准",
    "STJLR.07.5004": "变速器控制模块标定和仪表XCP通信协议工程标准",
    "STJLR.18.019": "CAN网络管理部件规范",
    "STJLR.18.039": "供应商软件开发设计流程要求",
    "STJLR.18.040": "供应商软件开发设计稳健性（TGW）",
    "STJLR.18.041": "制造诊断实施要求补充",
    "STJLR.18.082": "AUTOSAR通信（COM）工程标准",
    "STJLR.18.089": "Bootloader部件工程标准",
    "STJLR.18.090": "网络通信软件分区和内存管理工程标准",
    "STJLR.18.106": "诊断服务实施要求",
    "STJLR.18.107": "诊断数据实施要求",
    "STJLR.18.108": "软件交付格式",
    "STJLR.18.109": "软件数据压缩标准",
    "STJLR.18.124": "软件下载顺序标准",
    "STJLR.18.184": "以太网交换机软件下载工程标准",
    "STJLR.18.188": "开放诊断数据交换（ODX）编写规范",
    "STJLR.18.199": "空中软件升级（SOTA）：目标ECU兼容性",
    "STJLR.18.241": "48V及更高电压子系统部件电子设计规范",
    "STJLR.18.250": "软件质量要求",
    "TPJLR.00.076": "整车电气系统交互（含信息娱乐）评价",
    "TPJLR.00.479": "诊断功能试验规范",
    "TPJLR.18.002": "车辆静态电流测量试验程序",
    "TPJLR.18.068": "Bootloader部件设备级工程试验程序",
    "TPJLR.18.070": "AUTOSAR通信和网络层软件评审",
    "TPJLR.18.144": "电气负载电压和电流数据测量程序",
    "TPJLR.18.179": "ECU软件签名试验程序",
    "TPJLR.18.181": "安全软件开发工程试验程序",
    "TPJLR.18.199": "空中软件升级（SOTA）：SOTA车辆ECU兼容性试验程序",
    "STJLR.01.5021": "镀铬和亮饰件外观质量",
    "STJLR.01.5041": "车辆安全玻璃外观质量",
    "STJLR.01.5042": "内饰漆面装饰木饰件外观质量",
    "STJLR.01.5043": "漆面碳纤维部件外观质量",
    "STJLR.18.031": "质量因素（QF）",
    "TPJLR.00.003": "工作声音质量（整车级）试验程序",
    "TPJLR.00.344": "工作声音质量（部件级）",
    "TPJLR.01.229": "仪表板和中控台最大挠度及品质感",
    "TPJLR.18.026": "质量因素评审试验程序",
    "STJLR.00.5083": "能源管理标准",
    "STJLR.00.5107": "人机界面（HMI）设计策略：座舱显示",
    "STJLR.00.5164": "车辆区域定义",
    "STJLR.01.5022": "外饰可见分缝",
    "STJLR.01.5027": "外饰外露紧固件",
    "STJLR.01.5029": "外饰可见点焊",
    "STJLR.01.5030": "外饰可见注射点",
    "STJLR.01.5031": "外饰透视性",
    "STJLR.18.009": "车载部件电子设计规范",
    "STJLR.18.047": "配置MIM和规则集约束",
    "STJLR.18.059": "系统安全工作说明",
    "STJLR.18.102": "FOSS部件使用要求",
    "STJLR.18.105": "JLR通用安全算法",
    "STJLR.18.146": "车载清单数据块",
    "STJLR.18.195": "网络安全卫生要求",
    "STJLR.18.215": "可靠性开发接口协议",
    "STJLR.18.223": "SOTA能力车辆ECU的SOTA兼容性",
    "STJLR.18.226": "ECU安全启动",
    "STJLR.18.259": "CCCM视频接口物理层规范",
    "STJLR.AD.5005": "部件条码分配",
    "TPJLR.00.068": "试验车辆准备和试车检查程序",
    "TPJLR.00.071": "整车系统稳健性",
    "TPJLR.00.087": "试验项目中螺栓连接装配和监控",
    "TPJLR.00.201": "车辆和属性客户评分系统",
    "TPJLR.00.252": "车辆开关内饰评审方法",
    "TPJLR.01.119": "仪表板和中控台关键寿命试验",
    "TPJLR.01.273": "CAE准静态载荷施加程序",
    "TPJLR.15.054": "显示器运输保护试验程序",
    "TPJLR.18.036": "配置MIM和规则集约束试验程序",
    "TPJLR.18.139": "结露试验",
    "TPJLR.18.182": "加密/认证",
    "TPJLR.18.184": "运行安全加固",
    "TPJLR.18.186": "安全硬件开发试验程序",
    "TPJLR.18.200": "SOTA能力车辆ECU的SOTA兼容性试验程序",
    "TPJLR.18.202": "ECU安全启动工程试验程序",
    "TPJLR.18.242": "连接器壳体分离检查（CAD）",
    "TPJLR.18.243": "连接器壳体配合检查及NVH/ME批准（部件/台架）",
    "TPJLR.18.244": "连接器壳体公差叠加分离检查（CAD/图纸）",
}


@dataclass
class SourceRecord:
    category: str
    standard_code: str
    title: str
    title_zh: str
    title_en: str
    version: str
    date: str
    document_kind: str
    file_name: str
    relative_path: str
    absolute_path: str
    extension: str
    extraction_note: str


@dataclass
class StandardRecord:
    category: str
    standard_code: str
    title: str
    title_zh: str
    title_en: str
    version: str
    date: str
    document_kind: str
    source_count: int
    source_files: str
    source_paths: list[str]
    extraction_note: str


STOP_LABELS = (
    "Procedure:",
    "Standard:",
    "Issue:",
    "Date:",
    "Page:",
    "Printed copies",
    "Jaguar Land Rover",
    "A INDEX",
    "A Index",
    "B SCOPE",
    "Foreword",
    "Table of Contents",
)


def normalize_spaces(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = re.sub(r"[ \t\r\f\v]+", " ", value)
    value = re.sub(r"\n[ \t]*", "\n", value)
    return value.strip()


def compact_line(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip(" -_\t\r\n")


def extract_pdf_text(path: Path, page_limit: int = 2) -> tuple[str, str]:
    try:
        reader = PdfReader(str(path))
        chunks: list[str] = []
        for page in reader.pages[:page_limit]:
            chunks.append(page.extract_text() or "")
        text = normalize_spaces("\n".join(chunks))
        if text:
            return text, "pdf text"
        return "", "pdf text empty"
    except Exception as exc:  # pragma: no cover - diagnostic output only
        return "", f"pdf read error: {type(exc).__name__}: {exc}"


def extract_docx_text(path: Path, paragraph_limit: int = 80) -> tuple[str, str]:
    try:
        doc = Document(str(path))
        paragraphs: list[str] = []
        seen: set[str] = set()

        def append_text(value: str) -> None:
            text = value.strip()
            if text and text not in seen:
                seen.add(text)
                paragraphs.append(text)

        for section in doc.sections:
            for part in (section.header, section.first_page_header, section.even_page_header):
                for paragraph in part.paragraphs:
                    append_text(paragraph.text)
                for table in part.tables:
                    for row in table.rows:
                        append_text(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))

        for paragraph in doc.paragraphs[:paragraph_limit]:
            append_text(paragraph.text)
        for table in doc.tables[:5]:
            for row in table.rows[:20]:
                append_text(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))

        for xml_text in extract_docx_xml_text(path):
            append_text(xml_text)

        text = normalize_spaces("\n".join(paragraphs))
        if text:
            return text, "docx text"
        return "", "docx text empty"
    except Exception as exc:  # pragma: no cover - diagnostic output only
        return "", f"docx read error: {type(exc).__name__}: {exc}"


def extract_docx_xml_text(path: Path) -> list[str]:
    texts: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            names = sorted(
                name
                for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            )
            names = sorted(names, key=lambda name: (0 if "/header" in name else 1, name))
            for name in names:
                xml = archive.read(name)
                root = ET.fromstring(xml)
                for paragraph in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                    chunks = [
                        node.text.strip()
                        for node in paragraph.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                        if node.text and node.text.strip()
                    ]
                    if chunks:
                        texts.append(" ".join(chunks))
    except Exception:
        return []
    return texts


def extract_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    return "", "unsupported extension"


def category_for(path: Path) -> str:
    rel = path.relative_to(SOURCE_ROOT)
    parts = rel.parts[:-1]
    if not parts:
        return "Root"
    if parts[0] == "STJLR" and len(parts) > 1:
        return "/".join(parts[:2])
    return parts[0]


def classify_business_category(standard_code: str, title: str, relative_path: str = "") -> str:
    code = standard_code.upper()
    haystack = f"{standard_code} {title} {relative_path}".lower()

    if code.startswith("CADJLR") or "cadjlr" in haystack or "cad--requirement" in haystack:
        return "CAD/工程图纸"
    if re.search(r"\b(noise|sound quality|working sound|operating sound|squeak|rattle|buzz)\b", haystack):
        return "噪声测试"
    if "nvh" in haystack and not re.search(r"\b(connector|housing|mating|separation)\b", haystack):
        return "噪声测试"
    if re.search(r"\b(connector|housing|mating|separation|nvh|buck|bolted joint|bolted joints|deflection|load applicator)\b", haystack):
        return "结构/连接器/NVH"
    if "emc" in haystack or code.startswith("JLR-EMC"):
        return "EMC/电磁兼容"
    if re.search(r"\b(voc|aldehyde|odour|odor|air quality|toxproof|restricted substance|substance|fogging)\b", haystack):
        return "VOC/气味/禁限物质"
    if (
        "材料测试标准" in relative_path
        or code.startswith("TPJLR.52")
        or code.startswith("STJLR.50")
        or code.startswith("STJLR.51")
    ):
        return "材料性能测试"
    if (
        code in {"2048-EE", "TPJLR.18.125", "STJLR.18.233"}
        or re.search(
            r"\b(environment|environmental|thermal|temperature|humidity|climatic|climate|vibration|shock|durability|reliability|water ingress|ingress|dust|salt|corrosion|squeak|rattle|shaker|lifetime|life cycle|weather)\b",
            haystack,
        )
    ):
        return "环境可靠性测试"
    if re.search(
        r"\b(hud|head-up|head up|display|optical|windscreen|chromaticity|colour|color|reflectance|uniformity|distortion|zenbara|touch|field of view|fields of view)\b",
        haystack,
    ):
        return "光学/HUD/显示性能"
    if re.search(
        r"\b(material|plastic|polymer|leather|textile|adhesive|coating|paint|mould|mold|moldflow|fogging|flammability|substance)\b",
        haystack,
    ):
        return "材料性能测试"
    if re.search(r"\b(hmi|human machine|cockpit|switchgear|interior review|driver information|console|instrument panel)\b", haystack):
        return "HMI/开关/座舱"
    if re.search(r"\b(software|diagnostic|can|autosar|communication|bootloader|network|ethernet|sota|security|secure|cyber|odx|encryption|authentication|foss|algorithm)\b", haystack):
        return "软件/通信/网络安全"
    if re.search(r"\b(electrical|electronic|electronics|voltage|current|quiescent|load|loads|terminal|ecu|module|device level|48 v|48v|higher voltage)\b", haystack):
        if re.search(r"\b(test|measurement|procedure|evaluation)\b", haystack):
            return "电气/电子测试"
        return "电气/电子设计"
    if re.search(
        r"\b(appearance|cosmetic|quality|qf|feel|sound quality|customer rating|veneers|glass|chrome|bright work)\b",
        haystack,
    ):
        return "外观/感知质量"
    if re.search(
        r"\b(vehicle preparation|shakedown|whole vehicle|assembly|monitoring|test vehicle|systems robustness|attribute customer|method for conducting)\b",
        haystack,
    ):
        return "整车/装配/试验准备"
    if re.search(r"\b(process|management|requirements|standard requirements|statement of work|interface agreement|barcode|configuration mim|ruleset constraints|inventory|common security)\b", haystack):
        return "管理/流程/通用规范"
    return "其他参考"


def translate_title_to_chinese(standard_code: str, title: str) -> str:
    if standard_code in TITLE_ZH_BY_CODE:
        return TITLE_ZH_BY_CODE[standard_code]
    if re.search(r"[\u4e00-\u9fff]", title):
        return title
    return f"标准名称待确认：{title}"


def filename_stem_without_notes(path: Path) -> str:
    stem = path.stem
    stem = re.sub(r"\s*\([0-9]+\)$", "", stem)
    stem = re.sub(r"\s+iss(?:ue)?\s*[0-9.]+$", "", stem, flags=re.I)
    stem = re.sub(r"\s+Iss\s*[0-9.]+$", "", stem)
    stem = re.sub(r"_?v[0-9.]+(?:_Amendment_[0-9.]+)?$", "", stem, flags=re.I)
    return compact_line(stem)


def normalize_standard_code(value: str) -> str:
    value = compact_line(value)
    value = value.replace("_", "-")

    patterns = [
        (r"\b(TPJLR|STJLR|CADJLR)[\s.\-]*(\d{2})[\s.\-]*(\d{2,4})\b", "{0}.{1}.{2}"),
        (r"\b(STJLR)[\s.\-]*(AD)[\s.\-]*(\d{4})\b", "{0}.{1}.{2}"),
        (r"\b(JLR-HTG-\d+)\b", "{0}"),
        (r"\b(JLR-EMC-(?:CP|CS))\b", "{0}"),
        (r"\b(JLR-EMC-\d{4}-\d{4})\b", "{0}"),
        (r"\b(\d{4}-EE)\b", "{0}"),
    ]

    for pattern, fmt in patterns:
        match = re.search(pattern, value, flags=re.I)
        if not match:
            continue
        groups = [group.upper() for group in match.groups()]
        return fmt.format(*groups)

    return ""


def extract_code(text: str, path: Path) -> str:
    document_id = extract_label_value(text, "Document ID")
    if document_id:
        code = normalize_standard_code(document_id)
        if code:
            return code

    filename_code = normalize_standard_code(path.stem)
    if filename_code:
        return filename_code

    header_patterns = [
        r"\b(?:Procedure|Standard)\s*:\s*([A-Z0-9][A-Z0-9.\-\s]{3,40})",
        r"\b(JLR-HTG-\d+)\b",
        r"\b(JLR-EMC-(?:CP|CS|[0-9]{4}-[0-9]{4}))\b",
        r"\b(2048-EE)\b",
        r"\b(?:TPJLR|STJLR|CADJLR)[\s.\-]*\d{2}[\s.\-]*\d{3,4}\b",
        r"\bSTJLR[\s.\-]*AD[\s.\-]*\d{4}\b",
    ]
    for pattern in header_patterns:
        match = re.search(pattern, text, flags=re.I)
        if match:
            candidate = match.group(1) if match.lastindex else match.group(0)
            code = normalize_standard_code(candidate)
            if code:
                return code

    return filename_stem_without_notes(path)


def extract_label_value(text: str, label: str) -> str:
    lines = [compact_line(line) for line in text.splitlines() if compact_line(line)]
    for index, line in enumerate(lines):
        match = re.match(rf"^{re.escape(label)}\s*:\s*(.*)$", line, flags=re.I)
        if not match:
            continue
        values = [match.group(1).strip()]
        for following in lines[index + 1 : index + 4]:
            if any(following.startswith(stop) for stop in STOP_LABELS):
                break
            if re.match(r"^[A-Z](?:\s+[A-Z])?\s+[A-Z][A-Z\s]{2,}$", following):
                break
            values.append(following)
        return compact_line(" ".join(value for value in values if value))
    for line in lines:
        match = re.search(rf"\b{re.escape(label)}\s*:\s*(.+)$", line, flags=re.I)
        if not match:
            continue
        value = match.group(1)
        stop_pattern = r"\s+(?:Procedure|Standard|Issue|Date|Page|Title)\s*:"
        stop = re.search(stop_pattern, value, flags=re.I)
        if stop:
            value = value[: stop.start()]
        return compact_line(value)
    return ""


def extract_title(text: str, path: Path) -> str:
    title = extract_label_value(text, "Title")
    if title:
        return title

    if "JLR-EMC-CP" in path.name:
        return "EMC Compliance Procedure for Electrical/Electronic Components and Subsystems"
    if "JLR-EMC-CS" in path.name:
        return "EMC Component Standard"
    if path.name.startswith("2048-EE"):
        return "Environmental Test Selection Matrix"

    lines = [compact_line(line) for line in text.splitlines() if compact_line(line)]
    for index, line in enumerate(lines):
        if not re.search(r"(Engineering Standard|Test Procedure|CADD Standard|Engineering Procedure)", line, flags=re.I):
            continue
        values: list[str] = []
        for following in lines[index + 1 : index + 6]:
            if any(following.startswith(stop) for stop in STOP_LABELS):
                break
            if re.match(r"^(Procedure|Standard)\s*:", following, flags=re.I):
                break
            values.append(following)
        fallback = compact_line(" ".join(values))
        if fallback:
            return fallback

    stem = filename_stem_without_notes(path)
    stem = re.sub(r"^(TPJLR|STJLR|CADJLR)[\s.\-]*(?:AD[\s.\-]*)?\d{2,3}[\s.\-]*\d{3,4}\s*[-–—]?\s*", "", stem, flags=re.I)
    return compact_line(stem)


def extract_version_from_filename(path: Path) -> str:
    name = path.stem
    patterns = [
        r"\bv([0-9]+(?:\.[0-9]+)*)\b",
        r"\bAmendment[_\s-]*([0-9]+(?:\.[0-9]+)*)\b",
        r"\bIss(?:ue)?[_\s-]*([0-9]+(?:\.[0-9]+)*)\b",
        r"\biss([0-9]+(?:\.[0-9]+)*)\b",
        r"\bDraft\b",
    ]
    hits: list[str] = []
    for pattern in patterns:
        match = re.search(pattern, name, flags=re.I)
        if not match:
            continue
        if pattern == r"\bDraft\b":
            hits.append("Draft")
        elif pattern.lower().startswith(r"\bv"):
            hits.append(f"v{match.group(1)}")
        elif "amendment" in pattern.lower():
            hits.append(f"Amendment {match.group(1)}")
        else:
            hits.append(f"Issue {match.group(1)}")
    return " / ".join(dict.fromkeys(hits))


def extract_version(text: str, path: Path) -> str:
    patterns = [
        r"\bIssue\s*:\s*([A-Za-z0-9.]+)",
        r"\bIssue\s+([0-9]+(?:\.[0-9]+)*)\b",
        r"\bVersion\s*:\s*([A-Za-z0-9.]+)",
        r"\bVersion\s+([0-9]+(?:\.[0-9]+)*)\b",
        r"\bv([0-9]+(?:\.[0-9]+)*)\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.I)
        if match:
            value = match.group(1).strip()
            if pattern.startswith(r"\bv"):
                return f"v{value}"
            return f"Issue {value}" if pattern.lower().startswith(r"\bissue") else value

    filename_version = extract_version_from_filename(path)
    if filename_version:
        return filename_version
    return ""


def extract_date(text: str) -> str:
    date = extract_label_value(text, "Date")
    if date:
        return date

    patterns = [
        r"\b([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})\b",
        r"\b([0-9]{1,2}/[0-9]{1,2}/[0-9]{2,4})\b",
        r"\b([0-9]{4}-[0-9]{2}-[0-9]{2})\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1)
    return ""


def extract_document_kind(text: str, path: Path) -> str:
    if "CADD Standard" in text or path.name.upper().startswith("CADJLR"):
        return "CADD Standard"
    if "Engineering Standard" in text or path.name.upper().startswith("STJLR"):
        return "Engineering Standard"
    if "Test Procedure" in text or path.name.upper().startswith("TPJLR"):
        return "Test Procedure"
    if "Engineering Procedure" in text:
        return "Engineering Procedure"
    if "Environmental Test Selection Matrix" in path.name:
        return "Selection Matrix"
    return ""


def iter_source_files() -> Iterable[Path]:
    extensions = {".pdf", ".docx"}
    for path in sorted(SOURCE_ROOT.rglob("*")):
        if path.is_file() and path.suffix.lower() in extensions:
            yield path


def record_for(path: Path) -> SourceRecord:
    text, note = extract_text(path)
    code = extract_code(text, path)
    title = extract_title(text, path)
    title_zh = translate_title_to_chinese(code, title)
    version = extract_version(text, path)
    date = extract_date(text)
    kind = extract_document_kind(text, path)
    relative_path = str(path.relative_to(SOURCE_ROOT))
    return SourceRecord(
        category=classify_business_category(code, title, relative_path),
        standard_code=code,
        title=title,
        title_zh=title_zh,
        title_en=title,
        version=version,
        date=date,
        document_kind=kind,
        file_name=path.name,
        relative_path=relative_path,
        absolute_path=str(path),
        extension=path.suffix.lower().lstrip("."),
        extraction_note=note,
    )


def choose_best(records: list[SourceRecord]) -> SourceRecord:
    def score(record: SourceRecord) -> tuple[int, int, int, int]:
        return (
            1 if record.version else 0,
            1 if record.title else 0,
            1 if record.extension == "pdf" else 0,
            -len(record.relative_path),
        )

    return sorted(records, key=score, reverse=True)[0]


def build_standard_records(source_records: list[SourceRecord]) -> list[StandardRecord]:
    grouped: dict[str, list[SourceRecord]] = defaultdict(list)
    for record in source_records:
        grouped[record.standard_code or record.file_name].append(record)

    standards: list[StandardRecord] = []
    for code, records in grouped.items():
        best = choose_best(records)
        source_paths = sorted(record.relative_path for record in records)
        source_files = "; ".join(sorted(record.file_name for record in records))
        notes = "; ".join(sorted(set(record.extraction_note for record in records if record.extraction_note)))
        standards.append(
            StandardRecord(
                category=best.category,
                standard_code=code,
                title=best.title,
                title_zh=best.title_zh,
                title_en=best.title_en,
                version=best.version,
                date=best.date,
                document_kind=best.document_kind,
                source_count=len(records),
                source_files=source_files,
                source_paths=source_paths,
                extraction_note=notes,
            )
        )

    return sorted(standards, key=lambda item: (CATEGORY_ORDER.get(item.category, 99), item.standard_code, item.title))


def markdown_table(rows: list[StandardRecord]) -> str:
    lines = [
        "# JLR standards - name and version",
        "",
        f"- Source folder: `{SOURCE_ROOT}`",
        f"- Unique standards: {len(rows)}",
        "",
        "| No. | Category | Standard | Chinese Title | English Title | Version | Date |",
        "| ---: | --- | --- | --- | --- | --- | --- |",
    ]
    for index, row in enumerate(rows, 1):
        values = [
            str(index),
            row.category,
            row.standard_code,
            row.title_zh,
            row.title_en,
            row.version or "TBD",
            row.date,
        ]
        escaped = [value.replace("|", "\\|") for value in values]
        lines.append("| " + " | ".join(escaped) + " |")
    return "\n".join(lines) + "\n"


def main() -> int:
    if not SOURCE_ROOT.exists():
        print(f"Source folder not found: {SOURCE_ROOT}", file=sys.stderr)
        return 1

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    source_records = [record_for(path) for path in iter_source_files()]
    standard_records = build_standard_records(source_records)
    payload = {
        "source_root": str(SOURCE_ROOT),
        "source_count": len(source_records),
        "unique_count": len(standard_records),
        "sources": [asdict(record) for record in source_records],
        "standards": [asdict(record) for record in standard_records],
    }
    OUTPUT_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    OUTPUT_MD.write_text(markdown_table(standard_records), encoding="utf-8")
    print(json.dumps({k: payload[k] for k in ("source_root", "source_count", "unique_count")}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
