from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from scripts.extract_jlr_standards import (
    classify_business_category,
    extract_code,
    extract_docx_text,
    extract_title,
    normalize_standard_code,
    translate_title_to_chinese,
)


class ExtractJlrStandardsTest(unittest.TestCase):
    def test_classifies_standards_by_business_domain(self) -> None:
        cases = [
            ("TPJLR.18.125", "Electrical and Electronic Component Environmental Compatibility Test", "", "环境可靠性测试"),
            ("TPJLR.52.656", "Resistance to Creaking and Rubbing", "材料测试标准/TPJLR-52-656.pdf", "材料性能测试"),
            ("TPJLR.52.458", "Determination and Assessment of Odour from Interior Trim Materials", "TPJLR/TPJLR-52-458.pdf", "VOC/气味/禁限物质"),
            ("JLR-EMC-CP", "EMC Compliance Procedure", "JLR-EMC-CP_v1.5.pdf", "EMC/电磁兼容"),
            ("CADJLR.02.03", "Engineering Content - Standard Notes", "STJLR/CAD--requirement/CADJLR.02.03.pdf", "CAD/工程图纸"),
            ("STJLR.13.5009", "Head-Up Display Optical Performance Standard", "STJLR/STJLR.13.5009.pdf", "光学/HUD/显示性能"),
            ("STJLR.18.199", "Software Over the Air (SOTA): Target ECU Compatibility", "STJLR/STJLR-18-199.pdf", "软件/通信/网络安全"),
            ("STJLR.18.009", "Design Practices for Electronics in Vehicle Components", "STJLR/STJLR-18-009.docx", "电气/电子设计"),
            ("TPJLR.00.252", "Method for Conducting Interior Review of Vehicle Switchgear", "TPJLR/TPJLR-00-252.pdf", "HMI/开关/座舱"),
            ("TPJLR.00.068", "Test Vehicle Preparation and Shakedown", "TPJLR/TPJLR-00-068.pdf", "整车/装配/试验准备"),
            ("TPJLR.00.029", "Standard Vehicle Noise and Vibration Measurement Positions", "TPJLR/TPJLR-00-029.pdf", "噪声测试"),
            ("TPJLR.00.187", "Component or System Squeak and Rattle Shaker Test", "TPJLR/TPJLR-00-187.pdf", "噪声测试"),
            ("TPJLR.00.003", "Working Sound Quality (Whole Vehicle Level) Test Procedure", "TPJLR/TPJLR-00-003.pdf", "噪声测试"),
        ]

        for code, title, relative_path, expected in cases:
            with self.subTest(code=code):
                self.assertEqual(classify_business_category(code, title, relative_path), expected)

    def test_docx_extraction_includes_header_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "header-title.docx"
            document = Document()
            document.sections[0].header.paragraphs[0].text = "Header Standard Title"
            document.add_paragraph("Body Scope")
            document.save(docx_path)

            text, note = extract_docx_text(docx_path)

            self.assertEqual(note, "docx text")
            self.assertIn("Header Standard Title", text)

    def test_normalizes_short_cadjlr_codes(self) -> None:
        self.assertEqual(normalize_standard_code("CADJLR.02.03 Iss 3"), "CADJLR.02.03")

    def test_prefers_document_id_for_environmental_test_matrix(self) -> None:
        text = "Document ID:\nJLR-HTG-564459 issue 3\nEnvironmental Test Selection Matrix"

        self.assertEqual(
            extract_code(text, Path("2048-EE Environmental Test Selection Matrix.pdf")),
            "JLR-HTG-564459",
        )

    def test_translates_known_standard_titles_to_chinese(self) -> None:
        self.assertEqual(
            translate_title_to_chinese("STJLR.18.009", "Design Practices for Electronics in Vehicle Components"),
            "车载部件电子设计规范",
        )
        self.assertEqual(
            translate_title_to_chinese("TPJLR.18.125", "Electrical and Electronic Component Environmental Compatibility Test"),
            "电气和电子部件环境兼容性试验",
        )

    def test_extracts_title_embedded_after_header_metadata(self) -> None:
        text = (
            "Index\n"
            "Standard: STJLR.18.009 Issue: 9 Date: 15 April 2019 Page: 10 of 51 "
            "Jaguar Land Rover Limited - Engineering Standard "
            "Title: Design Practices for Electronics in Vehicle Components\n"
            "Scope\n"
        )

        self.assertEqual(
            extract_title(text, Path("STJLR-18-009-电气.docx")),
            "Design Practices for Electronics in Vehicle Components",
        )


if __name__ == "__main__":
    unittest.main()
